from django.shortcuts import render, redirect
from django.http import JsonResponse
from django.views.generic import ListView
from django.views.decorators.csrf import csrf_protect
from django.views.decorators.http import require_http_methods
from django.utils.decorators import method_decorator
from django.contrib import messages
from django.views import View
import json

from .repositories import (
    CarreraRepository, AsignaturaRepository, ProgramaAnaliticoRepository,
    UnidadRepository, PreguntaRepository, OpcionRepository, PartidaRepository
)
from .supabase_client import get_supabase_client


# ============================================================================
# VISTAS PRINCIPALES - SUPABASE
# ============================================================================

class PartidaListView(ListView):
    """Lista de partidas con contador de unidades - Supabase"""
    template_name = 'partidas/lista.html'
    context_object_name = 'partidas'
    paginate_by = 50

    def get_queryset(self):
        partidas = PartidaRepository.list_all(limit=1000)
        for partida in partidas:
            # Obtener asignatura
            asignatura = AsignaturaRepository.get_by_id(partida['asignatura_id'])
            if asignatura:
                partida['asignatura'] = asignatura

                # Obtener carrera si existe
                if asignatura.get('carrera_id'):
                    carrera = CarreraRepository.get_by_id(asignatura['carrera_id'])
                    if carrera:
                        partida['asignatura']['carrera'] = carrera

                # Calcular unidades
                programas = ProgramaAnaliticoRepository.list_by_asignatura(
                    asignatura_id=asignatura['asignatura_id']
                )
                total_unidades = 0
                for programa in programas:
                    unidades = UnidadRepository.list_all(
                        limit=1000,
                        programa_analitico_id=programa['linea_educativa_id']
                    )
                    total_unidades += len(unidades)
                partida['unidades_count'] = total_unidades
            else:
                partida['unidades_count'] = 0
                partida['asignatura'] = None

        return partidas

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        # Datos para el modal de crear partida
        context['asignaturas'] = AsignaturaRepository.list_all(limit=1000)
        context['carreras'] = CarreraRepository.list_all(limit=1000)
        return context


class UnidadListView(ListView):
    """Lista de unidades con filtros - Supabase"""
    template_name = 'unidades/lista.html'
    context_object_name = 'unidades'
    paginate_by = 50

    def get_queryset(self):
        try:
            # Obtener filtro de la URL
            programa_analitico_id = self.request.GET.get('programa_analitico')

            # Obtener unidades según el filtro
            if programa_analitico_id:
                unidades = UnidadRepository.list_all(
                    limit=1000,
                    programa_analitico_id=int(programa_analitico_id)
                )
            else:
                unidades = UnidadRepository.list_all(limit=1000)

            # Enriquecer datos de cada unidad
            for unidad in unidades:
                try:
                    # Obtener programa analítico
                    programa = ProgramaAnaliticoRepository.get_by_id(unidad['programa_analitico_id'])
                    if programa:
                        unidad['programa_analitico'] = programa

                        # Obtener asignatura
                        asignatura = AsignaturaRepository.get_by_id(programa['asignatura_id'])
                        if asignatura:
                            unidad['programa_analitico']['asignatura'] = asignatura

                            # Obtener carrera si existe
                            if asignatura.get('carrera_id'):
                                carrera = CarreraRepository.get_by_id(asignatura['carrera_id'])
                                if carrera:
                                    unidad['programa_analitico']['asignatura']['carrera'] = carrera

                            # Obtener partida asociada a la asignatura
                            partidas = PartidaRepository.list_all(limit=1000, asignatura_id=asignatura['asignatura_id'])
                            if partidas:
                                # Tomar la primera partida encontrada (asumiendo que hay una por asignatura)
                                unidad['partida'] = partidas[0]
                            else:
                                unidad['partida'] = None
                except Exception as e:
                    # Si falla el enriquecimiento de una unidad, continuar con las demás
                    import logging
                    logger = logging.getLogger(__name__)
                    logger.warning(f"Error al enriquecer unidad {unidad.get('unidad_id', 'unknown')}: {e}")
                    continue

            return unidades
        except Exception as e:
            import logging
            logger = logging.getLogger(__name__)
            logger.error(f"Error al obtener unidades: {e}")
            # Retornar lista vacía en caso de error
            return []

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        try:
            # Datos para el filtro de programas analíticos
            context['programas_analiticos'] = ProgramaAnaliticoRepository.list_all(limit=1000)

            # Enriquecer programas con asignatura y carrera
            for programa in context['programas_analiticos']:
                try:
                    asignatura = AsignaturaRepository.get_by_id(programa['asignatura_id'])
                    if asignatura:
                        programa['asignatura'] = asignatura
                        if asignatura.get('carrera_id'):
                            carrera = CarreraRepository.get_by_id(asignatura['carrera_id'])
                            if carrera:
                                programa['asignatura']['carrera'] = carrera
                except Exception as e:
                    import logging
                    logger = logging.getLogger(__name__)
                    logger.warning(f"Error al enriquecer programa {programa.get('linea_educativa_id', 'unknown')}: {e}")
                    continue
        except Exception as e:
            import logging
            logger = logging.getLogger(__name__)
            logger.error(f"Error al obtener programas analíticos: {e}")
            context['programas_analiticos'] = []

        return context


class PreguntaListView(ListView):
    """Lista de preguntas con filtros - Supabase"""
    template_name = 'preguntas/lista.html'
    context_object_name = 'preguntas'
    paginate_by = 50

    def get_queryset(self):
        try:
            partida_id = self.request.GET.get('partida')
            if not partida_id:
                return []

            partida = PartidaRepository.get_by_id(int(partida_id))
            if not partida:
                return []

            asignatura = AsignaturaRepository.get_by_id(partida['asignatura_id'])
            if not asignatura:
                return []

            programas = ProgramaAnaliticoRepository.list_by_asignatura(
                asignatura_id=asignatura['asignatura_id']
            )

            programa_analitico_id = self.request.GET.get('programa_analitico')
            if programa_analitico_id:
                programas = [p for p in programas if p['linea_educativa_id'] == int(programa_analitico_id)]

            todas_unidades = []
            for programa in programas:
                try:
                    unidades = UnidadRepository.list_all(
                        limit=1000,
                        programa_analitico_id=programa['linea_educativa_id']
                    )
                    todas_unidades.extend(unidades)
                except Exception as e:
                    import logging
                    logger = logging.getLogger(__name__)
                    logger.warning(f"Error al obtener unidades para programa {programa.get('linea_educativa_id', 'unknown')}: {e}")
                    continue

            unidad_id = self.request.GET.get('unidad')
            if unidad_id:
                todas_unidades = [u for u in todas_unidades if u['unidad_id'] == int(unidad_id)]

            todas_preguntas = []
            for unidad in todas_unidades:
                try:
                    preguntas = PreguntaRepository.list_all(limit=1000, unidad_id=unidad['unidad_id'])
                    for pregunta in preguntas:
                        pregunta['unidad'] = unidad
                        programa_correspondiente = next(
                            (p for p in programas if p['linea_educativa_id'] == unidad['programa_analitico_id']),
                            None
                        )
                        if programa_correspondiente:
                            pregunta['programa_analitico'] = programa_correspondiente
                            pregunta['asignatura'] = asignatura
                        try:
                            pregunta['opciones'] = OpcionRepository.list_all(limit=1000, pregunta_id=pregunta['pregunta_id'])
                        except Exception as e:
                            import logging
                            logger = logging.getLogger(__name__)
                            logger.warning(f"Error al obtener opciones para pregunta {pregunta.get('pregunta_id', 'unknown')}: {e}")
                            pregunta['opciones'] = []
                    todas_preguntas.extend(preguntas)
                except Exception as e:
                    import logging
                    logger = logging.getLogger(__name__)
                    logger.warning(f"Error al obtener preguntas para unidad {unidad.get('unidad_id', 'unknown')}: {e}")
                    continue

            todas_preguntas.sort(key=lambda x: int(x.get('numero', 0)))
            return todas_preguntas
        except Exception as e:
            import logging
            logger = logging.getLogger(__name__)
            logger.error(f"Error al obtener preguntas: {e}")
            return []

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        try:
            context['partidas'] = PartidaRepository.list_all(limit=1000)
        except Exception as e:
            import logging
            logger = logging.getLogger(__name__)
            logger.error(f"Error al obtener partidas: {e}")
            context['partidas'] = []
        
        try:
            context['unidades'] = UnidadRepository.list_all(limit=1000)
        except Exception as e:
            import logging
            logger = logging.getLogger(__name__)
            logger.error(f"Error al obtener unidades: {e}")
            context['unidades'] = []
        
        try:
            context['programas_analiticos'] = ProgramaAnaliticoRepository.list_all(limit=1000)
        except Exception as e:
            import logging
            logger = logging.getLogger(__name__)
            logger.error(f"Error al obtener programas analíticos: {e}")
            context['programas_analiticos'] = []
        
        return context


# ============================================================================
# APIs DE EDICIÓN INLINE - SUPABASE
# ============================================================================

@csrf_protect
@require_http_methods(["POST"])
def update_partida_api(request, partida_id):
    """API para actualizar nombre de partida - Supabase"""
    try:
        data = json.loads(request.body)
        if 'descripcion' in data:
            partida = PartidaRepository.update(int(partida_id), descripcion=data['descripcion'])
            return JsonResponse({
                'success': True,
                'message': 'Partida actualizada exitosamente',
                'partida': {
                    'id': partida['partida_id'],
                    'descripcion': partida['descripcion']
                }
            })
        else:
            return JsonResponse({'success': False, 'error': 'Descripción requerida'}, status=400)
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=400)


@csrf_protect
@require_http_methods(["POST"])
def update_unidad_api(request, unidad_id):
    """API para actualizar descripción de unidad - Supabase"""
    try:
        data = json.loads(request.body)
        if 'descripcion' in data:
            unidad = UnidadRepository.update(int(unidad_id), descripcion=data['descripcion'])
            return JsonResponse({
                'success': True,
                'message': 'Unidad actualizada exitosamente',
                'unidad': {
                    'id': unidad['unidad_id'],
                    'descripcion': unidad['descripcion'],
                    'numero_unidad': unidad['numero_unidad']
                }
            })
        else:
            return JsonResponse({'success': False, 'error': 'Descripción requerida'}, status=400)
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=400)


# ============================================================================
# APIs DE CREACIÓN - SUPABASE
# ============================================================================

@csrf_protect
@require_http_methods(["POST"])
def crear_carrera_ajax(request):
    """API para crear carrera - Supabase"""
    try:
        data = json.loads(request.body)
        descripcion = data.get('descripcion', '').strip()

        if not descripcion:
            return JsonResponse({'success': False, 'error': 'La descripción es requerida'}, status=400)

        carrera = CarreraRepository.create(descripcion=descripcion)
        return JsonResponse({
            'success': True,
            'message': 'Carrera creada exitosamente',
            'carrera': {
                'id': carrera['carrera_id'],
                'descripcion': carrera['descripcion']
            }
        })
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=400)


@csrf_protect
@require_http_methods(["POST"])
def crear_asignatura_ajax(request):
    """API para crear asignatura - Supabase"""
    try:
        data = json.loads(request.body)
        descripcion = data.get('descripcion', '').strip()
        carrera_id = data.get('carrera_id')

        if not descripcion:
            return JsonResponse({'success': False, 'error': 'La descripción es requerida'}, status=400)

        if not carrera_id:
            return JsonResponse({'success': False, 'error': 'La carrera es requerida'}, status=400)

        asignatura = AsignaturaRepository.create(
            descripcion=descripcion,
            carrera_id=int(carrera_id)
        )
        return JsonResponse({
            'success': True,
            'message': 'Asignatura creada exitosamente',
            'asignatura': {
                'id': asignatura['asignatura_id'],
                'descripcion': asignatura['descripcion'],
                'carrera_id': asignatura['carrera_id']
            }
        })
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=400)


# ============================================================================
# APIs DE FILTROS DINÁMICOS - SUPABASE
# ============================================================================

def get_programas_analiticos(request):
    """API para obtener programas analíticos por partida o por asignatura (Supabase)."""
    try:
        partida_id = request.GET.get('partida_id')
        asignatura_id = request.GET.get('asignatura_id')

        # 1) Si viene partida_id, obtenemos la asignatura desde la partida
        if partida_id:
            try:
                pid = int(partida_id)
            except ValueError:
                return JsonResponse({'error': 'partida_id inválido'}, status=400)

            try:
                partida = PartidaRepository.get_by_id(pid)
                print(f"DEBUG: partida encontrada={partida}")
            except Exception as e:
                print(f"DEBUG: Error al obtener partida: {e}")
                return JsonResponse({'error': f'Error al obtener partida: {e}'}, status=400)

            if not partida:
                print("DEBUG: Partida no encontrada, devolviendo array vacío")
                return JsonResponse([], safe=False)

            # Si la partida existe, usamos su asignatura_id (a menos que ya venga uno explícito)
            asignatura_id = asignatura_id or partida.get('asignatura_id')
            print(f"DEBUG: asignatura_id de partida={asignatura_id}")

        # 2) Validación: necesitamos asignatura_id para listar programas
        if not asignatura_id:
            print("DEBUG: No hay asignatura_id, devolviendo array vacío")
            return JsonResponse([], safe=False)

        try:
            aid = int(asignatura_id)
        except ValueError:
            return JsonResponse({'error': 'asignatura_id inválido'}, status=400)

        # 3) Obtener programas por asignatura
        try:
            programas = ProgramaAnaliticoRepository.list_by_asignatura(asignatura_id=aid)
            print(f"DEBUG: programas encontrados={len(programas) if programas else 0}")
            return JsonResponse(programas or [], safe=False)
        except Exception as e:
            print(f"DEBUG: Error al obtener programas: {e}")
            return JsonResponse({'error': f'Error al obtener programas: {e}'}, status=400)

    except Exception as e:
        print(f"DEBUG: Error general en get_programas_analiticos: {e}")
        return JsonResponse({'error': str(e)}, status=500)


def get_unidades(request):
    """API para obtener unidades por programa analítico - Supabase"""
    try:
        programa_id = request.GET.get('programa_id')
        if not programa_id:
            return JsonResponse([], safe=False)

        unidades = UnidadRepository.list_all(
            limit=1000,
            programa_analitico_id=int(programa_id)
        )
        return JsonResponse(unidades, safe=False)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)


@csrf_protect
@require_http_methods(["POST"])
def update_pregunta_api(request, pregunta_id):
    """API para actualizar pregunta - Supabase"""
    try:
        data = json.loads(request.body)
        enunciado = data.get('enunciado', '').strip()
        explicacion = data.get('explicacion', '').strip()
        opciones = data.get('opciones', [])

        if not enunciado:
            return JsonResponse({'success': False, 'error': 'Enunciado requerido'}, status=400)

        # Actualizar enunciado y explicación de la pregunta
        pregunta = PreguntaRepository.update(int(pregunta_id), enunciado=enunciado, explicacion=explicacion)
        if not pregunta:
            return JsonResponse({'success': False, 'error': 'Error al actualizar pregunta'}, status=400)

        # Actualizar opciones
        for opcion_data in opciones:
            opcion_id = opcion_data.get('id')
            texto = opcion_data.get('texto', '').strip()
            es_correcta = opcion_data.get('es_correcta', False)

            if opcion_id and opcion_id != 'new':
                # Actualizar opción existente
                OpcionRepository.update(
                    int(opcion_id),
                    opcion=texto,
                    es_correcta=es_correcta
                )
            elif opcion_id == 'new' and texto:
                # Crear nueva opción
                OpcionRepository.create(
                    opcion=texto,
                    es_correcta=es_correcta,
                    pregunta_id=int(pregunta_id)
                )

        return JsonResponse({
            'success': True,
            'message': 'Pregunta actualizada exitosamente'
        })
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=400)


@csrf_protect
@require_http_methods(["DELETE"])
def delete_opcion_api(request, opcion_id):
    """API para eliminar opción - Supabase"""
    try:
        success = OpcionRepository.delete(int(opcion_id))
        if success:
            return JsonResponse({'success': True, 'message': 'Opción eliminada exitosamente'})
        else:
            return JsonResponse({'success': False, 'error': 'Error al eliminar opción'}, status=400)
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=400)


def obtener_prompt(request):
    """API para obtener prompt de generación de preguntas - Supabase"""
    try:
        partida_id = request.GET.get('partida')
        programa_id = request.GET.get('programa_analitico')
        unidad_id = request.GET.get('unidad')

        if not partida_id:
            return JsonResponse({'success': False, 'error': 'Partida requerida'}, status=400)

        # Obtener datos de la partida
        partida = PartidaRepository.get_by_id(int(partida_id))
        if not partida:
            return JsonResponse({'success': False, 'error': 'Partida no encontrada'}, status=404)

        # Obtener asignatura
        asignatura = AsignaturaRepository.get_by_id(partida['asignatura_id'])
        if not asignatura:
            return JsonResponse({'success': False, 'error': 'Asignatura no encontrada'}, status=404)

        # Obtener carrera
        carrera = None
        if asignatura.get('carrera_id'):
            carrera = CarreraRepository.get_by_id(asignatura['carrera_id'])

        # Obtener programas analíticos
        programas = ProgramaAnaliticoRepository.list_by_asignatura(
            asignatura_id=asignatura['asignatura_id']
        )

        # Obtener unidades y preguntas
        unidades_data = []
        unidad_actual = None  # Para almacenar la unidad actual si se filtra

        for programa in programas:
            unidades = UnidadRepository.list_all(
                limit=1000,
                programa_analitico_id=programa['linea_educativa_id']
            )

            for unidad in unidades:
                # Filtrar por unidad específica si se proporciona
                if unidad_id and unidad['unidad_id'] != int(unidad_id):
                    continue

                # Guardar la unidad actual si se está filtrando por unidad específica
                if unidad_id and unidad['unidad_id'] == int(unidad_id):
                    unidad_actual = {
                        'numero': unidad['numero_unidad'],
                        'descripcion': unidad['descripcion']
                    }

                # Obtener preguntas de la unidad
                preguntas = PreguntaRepository.list_all(limit=1000, unidad_id=unidad['unidad_id'])

                unidad_info = {
                    'numero': unidad['numero_unidad'],
                    'descripcion': unidad['descripcion'],
                    'preguntas': []
                }

                for pregunta in preguntas:
                    # Obtener opciones de la pregunta
                    opciones = OpcionRepository.list_all(limit=1000, pregunta_id=pregunta['pregunta_id'])

                    pregunta_info = {
                        'numero': pregunta['numero'],
                        'enunciado': pregunta['enunciado'],
                        'opciones': []
                    }

                    for opcion in opciones:
                        pregunta_info['opciones'].append({
                            'texto': opcion['opcion'],
                            'es_correcta': opcion['es_correcta']
                        })

                    unidad_info['preguntas'].append(pregunta_info)

                if unidad_info['preguntas']:  # Solo agregar si tiene preguntas
                    unidades_data.append(unidad_info)

        # Generar prompt con información de unidad actual
        prompt = generar_prompt_texto(
            partida, asignatura, carrera, programas, unidades_data, unidad_actual
        )

        return JsonResponse({
            'success': True,
            'prompt': prompt
        })

    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=400)


def descargar_google_docs(request):
    """API para descargar documento de Google Docs - Supabase"""
    try:
        from django.http import HttpResponse
        from docx import Document
        from docx.shared import Inches
        import io
        import logging

        logger = logging.getLogger(__name__)
        logger.info(f"Iniciando descarga de documento para partida: {request.GET.get('partida')}")

        partida_id = request.GET.get('partida')
        programa_id = request.GET.get('programa_analitico')
        unidad_id = request.GET.get('unidad')

        if not partida_id:
            logger.error("Partida no proporcionada")
            return JsonResponse({'success': False, 'error': 'Partida requerida'}, status=400)

        # Obtener datos (similar a obtener_prompt)
        partida = PartidaRepository.get_by_id(int(partida_id))
        if not partida:
            logger.error(f"Partida {partida_id} no encontrada")
            return JsonResponse({'success': False, 'error': 'Partida no encontrada'}, status=404)

        asignatura = AsignaturaRepository.get_by_id(partida['asignatura_id'])
        if not asignatura:
            logger.error(f"Asignatura {partida['asignatura_id']} no encontrada")
            return JsonResponse({'success': False, 'error': 'Asignatura no encontrada'}, status=404)

        carrera = None
        if asignatura.get('carrera_id'):
            carrera = CarreraRepository.get_by_id(asignatura['carrera_id'])

        # Obtener datos de preguntas con límites para evitar documentos muy grandes
        unidades_data = []
        programas = ProgramaAnaliticoRepository.list_by_asignatura(
            asignatura_id=asignatura['asignatura_id']
        )

        total_preguntas = 0
        max_preguntas = 500  # Límite para evitar documentos muy grandes

        for programa in programas:
            if total_preguntas >= max_preguntas:
                logger.warning(f"Límite de {max_preguntas} preguntas alcanzado")
                break
                
            unidades = UnidadRepository.list_all(
                limit=100,
                programa_analitico_id=programa['linea_educativa_id']
            )

            for unidad in unidades:
                if total_preguntas >= max_preguntas:
                    break
                    
                # Limitar preguntas por unidad
                preguntas = PreguntaRepository.list_all(limit=50, unidad_id=unidad['unidad_id'])

                unidad_info = {
                    'numero': unidad['numero_unidad'],
                    'descripcion': unidad['descripcion'],
                    'preguntas': []
                }

                # Ordenar preguntas por número
                preguntas_ordenadas = sorted(preguntas, key=lambda x: int(x.get('numero', 0)))
                
                for pregunta in preguntas_ordenadas:
                    if total_preguntas >= max_preguntas:
                        break
                        
                    opciones = OpcionRepository.list_all(limit=10, pregunta_id=pregunta['pregunta_id'])

                    pregunta_info = {
                        'numero': pregunta['numero'],
                        'enunciado': pregunta['enunciado'],
                        'explicacion': pregunta.get('explicacion', ''),
                        'opciones': []
                    }

                    for opcion in opciones:
                        pregunta_info['opciones'].append({
                            'texto': opcion['opcion'],
                            'es_correcta': opcion['es_correcta']
                        })

                    unidad_info['preguntas'].append(pregunta_info)
                    total_preguntas += 1

                if unidad_info['preguntas']:
                    unidades_data.append(unidad_info)

        logger.info(f"Generando documento con {total_preguntas} preguntas en {len(unidades_data)} unidades")

        # Verificar que hay datos para generar
        if not unidades_data:
            logger.warning("No se encontraron preguntas para generar el documento")
            return JsonResponse({'success': False, 'error': 'No se encontraron preguntas para generar el documento'}, status=404)

        # Generar documento Word
        doc = generar_documento_word(partida, asignatura, carrera, unidades_data)

        # Crear respuesta HTTP
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
        # Sanitizar nombre de archivo
        safe_filename = "".join(c for c in partida["descripcion"] if c.isalnum() or c in (' ', '-', '_')).rstrip()
        response['Content-Disposition'] = f'attachment; filename="preguntas_{safe_filename}.docx"'

        # Guardar documento en memoria
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)

        response.write(doc_buffer.getvalue())
        logger.info("Documento generado exitosamente")
        return response

    except Exception as e:
        logger.error(f"Error al generar documento: {str(e)}", exc_info=True)
        return JsonResponse({'success': False, 'error': f'Error interno: {str(e)}'}, status=500)


def extraer_contexto_por_unidad(contexto_completo, numero_unidad):
    """
    Extrae el contexto específico de una unidad del contexto completo del programa analítico.
    Maneja el formato real: Contexto general + Unidades con temas y subtemas.
    """
    if not contexto_completo or not numero_unidad:
        return contexto_completo
    
    # Convertir a string y limpiar
    contexto = str(contexto_completo).strip()
    numero_unidad = int(numero_unidad)
    
    # Buscar la sección "Unidades, temas y subtemas"
    seccion_unidades = "Unidades, temas y subtemas"
    pos_seccion = contexto.lower().find(seccion_unidades.lower())
    
    if pos_seccion == -1:
        # Si no encuentra la sección, buscar directamente la unidad
        return _buscar_unidad_directa(contexto, numero_unidad)
    
    # Extraer desde la sección de unidades hasta el final
    contexto_unidades = contexto[pos_seccion:].strip()
    
    # Buscar la unidad específica
    patrones_unidad = [
        f"Unidad {numero_unidad}:",
        f"Unidad {numero_unidad} ",
    ]
    
    inicio_unidad = -1
    for patron in patrones_unidad:
        pos = contexto_unidades.lower().find(patron.lower())
        if pos != -1:
            inicio_unidad = pos
            break
    
    if inicio_unidad == -1:
        # Si no encuentra la unidad específica, devolver todo el contexto
        return contexto
    
    # Buscar el final de la unidad (siguiente unidad o final del texto)
    siguiente_unidad = numero_unidad + 1
    fin_unidad = len(contexto_unidades)
    
    # Buscar la siguiente unidad
    patrones_siguiente = [
        f"Unidad {siguiente_unidad}:",
        f"Unidad {siguiente_unidad} ",
    ]
    
    for patron in patrones_siguiente:
        pos = contexto_unidades.lower().find(patron.lower(), inicio_unidad + 1)
        if pos != -1:
            fin_unidad = pos
            break
    
    # Extraer el contexto de la unidad
    contexto_unidad = contexto_unidades[inicio_unidad:fin_unidad].strip()
    
    # Si el contexto extraído es muy corto, devolver el contexto completo
    if len(contexto_unidad) < 50:
        return contexto
    
    # Construir el contexto completo: contexto general + unidad específica
    contexto_general = contexto[:pos_seccion].strip()
    
    # Buscar el título de la sección de unidades
    titulo_seccion = "Unidades, temas y subtemas (tal como en el programa)"
    if titulo_seccion.lower() in contexto.lower():
        # Incluir el título de la sección
        resultado = f"{contexto_general}\n\n{titulo_seccion}\n\n{contexto_unidad}"
    else:
        resultado = f"{contexto_general}\n\n{seccion_unidades}\n\n{contexto_unidad}"
    
    return resultado


def _buscar_unidad_directa(contexto, numero_unidad):
    """
    Busca la unidad directamente en el contexto sin sección específica.
    """
    patrones_unidad = [
        f"Unidad {numero_unidad}:",
        f"Unidad {numero_unidad} ",
        f"Tema {numero_unidad}:",
        f"Tema {numero_unidad} ",
        f"Capítulo {numero_unidad}:",
        f"Capítulo {numero_unidad} ",
    ]
    
    # Buscar el inicio de la unidad
    inicio_unidad = -1
    for patron in patrones_unidad:
        pos = contexto.lower().find(patron.lower())
        if pos != -1:
            inicio_unidad = pos
            break
    
    if inicio_unidad == -1:
        # Si no encuentra patrones específicos, buscar por número
        import re
        patron_numero = rf'\b{numero_unidad}\.?\s*[:\-]?\s*'
        match = re.search(patron_numero, contexto, re.IGNORECASE)
        if match:
            inicio_unidad = match.start()
    
    if inicio_unidad == -1:
        # Si no encuentra la unidad específica, devolver todo el contexto
        return contexto
    
    # Buscar el final de la unidad (siguiente unidad o final del texto)
    siguiente_unidad = numero_unidad + 1
    fin_unidad = len(contexto)
    
    # Buscar la siguiente unidad
    for siguiente_num in range(siguiente_unidad, siguiente_unidad + 3):  # Buscar hasta 3 unidades adelante
        patrones_siguiente = [
            f"Unidad {siguiente_num}:",
            f"Unidad {siguiente_num} ",
            f"Tema {siguiente_num}:",
            f"Tema {siguiente_num} ",
            f"Capítulo {siguiente_num}:",
            f"Capítulo {siguiente_num} ",
        ]
        
        for patron in patrones_siguiente:
            pos = contexto.lower().find(patron.lower(), inicio_unidad + 1)
            if pos != -1:
                fin_unidad = pos
                break
        
        if fin_unidad < len(contexto):
            break
    
    # Extraer el contexto de la unidad
    contexto_unidad = contexto[inicio_unidad:fin_unidad].strip()
    
    # Si el contexto extraído es muy corto, devolver el contexto completo
    if len(contexto_unidad) < 50:
        return contexto
    
    return contexto_unidad


def generar_prompt_texto(partida, asignatura, carrera, programas, unidades_data, unidad_actual=None):
    """Generar texto del prompt para IA con formato específico"""

    # Obtener programa analítico (usar el primero si hay varios)
    programa = programas[0] if programas else {}

    # Construir lista de unidades con número de preguntas y rangos de numeración
    unidades_lista = []
    pregunta_inicio = 1
    
    for unidad in unidades_data:
        num_preguntas = len(unidad['preguntas'])
        pregunta_fin = pregunta_inicio + num_preguntas - 1
        
        unidades_lista.append({
            "numero": unidad['numero'],
            "descripcion": unidad['descripcion'],
            "num_preguntas": num_preguntas,
            "rango_preguntas": f"{pregunta_inicio}-{pregunta_fin}"
        })
        
        pregunta_inicio = pregunta_fin + 1

    # Filtrar contexto por unidad si se especifica una unidad actual
    contexto_filtrado = programa.get('contexto', 'No especificado')
    if unidad_actual:
        contexto_filtrado = extraer_contexto_por_unidad(
            programa.get('contexto', ''), 
            unidad_actual['numero']
        )
        
        # Si no se pudo extraer contexto específico, usar el completo pero con advertencia
        if contexto_filtrado == programa.get('contexto', ''):
            contexto_filtrado = f"CONTEXTO COMPLETO (no se pudo filtrar por unidad {unidad_actual['numero']}):\n{contexto_filtrado}"
        else:
            contexto_filtrado = f"CONTEXTO DE LA UNIDAD {unidad_actual['numero']}:\n{contexto_filtrado}"


    # Generar el prompt con el formato específico
    prompt = f"""ROL
Eres un reclutador de docentes de universidad de élite. Debes evaluar un concurso observando cómo enseñan y cómo se desempeñan con preguntas claras, de dificultad baja–media, enfocadas en conceptos fundamentales y su aplicación directa.

ENTRADAS
- Carrera: {carrera['descripcion'] if carrera else 'No especificada'}
- Programa: {asignatura['descripcion']}
- Asignatura: {asignatura['descripcion']}
- Plan analítico (título): {programa.get('titulo', 'No especificado')}
- Contexto del plan: {contexto_filtrado}  ← aquí se listan los temas y subtemas por unidad; úsalos como base obligatoria

OBJETIVO
Para cada unidad, genera EXACTAMENTE el número de preguntas solicitado, tipo **caso breve de estudio (nivel universitario)**, con **una sola respuesta correcta** por pregunta. Cada caso debe alinearse explícitamente con **un tema/subtema** de esa unidad indicado en el **Contexto del plan**.

NUMERACIÓN DE PREGUNTAS:
- Las preguntas deben numerarse secuencialmente según el rango indicado para cada unidad
- Unidad 1: Preguntas 1-25 (o según el número de preguntas configurado)
- Unidad 2: Preguntas 26-50 (continuando desde donde terminó la unidad anterior)
- Unidad 3: Preguntas 51-75 (y así sucesivamente)
- Cada unidad debe generar preguntas en su rango específico sin solapamientos

IMPORTANTE - EXTRACCIÓN DE TEMAS/SUBTEMAS:
- Los temas y subtemas para cada pregunta DEBEN extraerse del "Contexto del plan" mostrado arriba
- Cada unidad tiene sus temas específicos listados en el contexto
- NO inventes temas, usa SOLO los que aparecen en el contexto para cada unidad
- Los temas/subtemas deben ser específicos de la unidad correspondiente

ESTRUCTURA DE CADA PREGUNTA (formato específico)
- **Narración del caso (4–5 líneas):** escenario realista y corto (laboratorio, aula, empresa, servicio público, clínica, etc.) que conecte con un **tema/subtema** de la unidad.
- **Pregunta directa (3 líneas):** compara, identifica o aplica un concepto clave del **tema/subtema**.
- **Opciones (A–D):** cuatro opciones, **solo una correcta**. Distractores plausibles, sin trucos rebuscados.
- **Respuesta correcta:** letra de la opción correcta.
- **Explicación (3-4 líneas):** razón breve de por qué es correcta y por qué las otras no.

REQUISITOS DE CALIDAD (EN TODO EL SET)
- Lenguaje claro y didáctico; evita tecnicismos innecesarios.
- **Cobertura balanceada:** usa todos los temas/subtemas del Contexto para esa unidad (si sobran preguntas, repite con enfoque distinto).
- **Dificultad global:** ~70% baja y ~30% media. No incluir preguntas de alta dificultad.
- **Enfoque conceptual:** una idea central por pregunta (definición, diferencia clave, identificación, aplicación directa).
- **Escenarios breves y variados**, pero simples; sin cálculos extensos ni información ambigua.
- **Distractores** basados en confusiones comunes; no uses redacciones capciosas.
- Longitud por ítem (sin metadatos): **60–120 palabras**.
- **Una sola** opción correcta; no repetir la respuesta con sinónimos en otras opciones.
- Referenciar explícitamente la **Unidad** y el **Tema/Subtema** en cada ítem.

INSTRUCCIONES ESPECÍFICAS PARA TEMAS/SUBTEMAS:
1. Revisa el "Contexto del plan" mostrado arriba
2. Identifica los temas y subtemas listados para cada unidad
3. Para cada pregunta, usa UNO de esos temas/subtemas específicos
4. NO inventes temas que no estén en el contexto
5. Cada unidad debe tener preguntas basadas en SUS temas específicos del contexto



FORMATO DE SALIDA (JSON específico)
Genera las preguntas en formato JSON con la siguiente estructura exacta:

{{
  "numero": [número de pregunta],
  "enunciado": "[Narración del caso + pregunta directa en un solo texto]",
  "explicacion": "[Explicación de 2-3 líneas sobre por qué es correcta la respuesta]",
  "opciones": [
    {{
      "texto": "[Texto de la opción A]",
      "correcta": true/false
    }},
    {{
      "texto": "[Texto de la opción B]", 
      "correcta": true/false
    }},
    {{
      "texto": "[Texto de la opción C]",
      "correcta": true/false
    }},
    {{
      "texto": "[Texto de la opción D]",
      "correcta": true/false
    }}
  ]
}}

EJEMPLO DE FORMATO CORRECTO:
{{
  "numero": 1,
  "enunciado": "En el diario 'Horizonte', las notas llegan tarde a edición porque los redactores escriben párrafos introductorios largos y dejan el dato clave para el tercer bloque. El editor quiere mejorar la comprensión en móviles y bajar la tasa de rebote en la web. Propone un taller express antes del cierre.\\n¿Qué decisión de redacción aplica mejor al caso para garantizar claridad inmediata?",
  "explicacion": "La pirámide invertida y un lead informativo priorizan lo esencial, mejoran escaneabilidad y tiempo de lectura. Citas largas, suspenso y primera persona no responden a la necesidad de claridad y rapidez.",
  "opciones": [
    {{
      "texto": "Usar lead informativo con 5W1H e invertir el orden (pirámide invertida)",
      "correcta": true
    }},
    {{
      "texto": "Iniciar con una cita literaria para enganchar emociones",
      "correcta": false
    }},
    {{
      "texto": "Dejar los datos clave para el cierre y 'suspenso'",
      "correcta": false
    }},
    {{
      "texto": "Escribir en primera persona para 'humanizar'",
      "correcta": false
    }}
  ]
}}

INSTRUCCIONES ESPECÍFICAS PARA EL FORMATO JSON:
1. Cada pregunta debe ser un objeto JSON válido
2. El campo "numero" debe corresponder al rango de la unidad
3. El campo "enunciado" debe incluir tanto la narración del caso como la pregunta directa
4. El campo "explicacion" debe ser una explicación clara de 2-3 líneas
5. El campo "opciones" debe ser un array con exactamente 4 opciones
6. Solo UNA opción debe tener "correcta": true, las otras tres deben tener "correcta": false
7. Genera todas las preguntas en un array JSON válido
8. Usa comillas dobles para strings, no comillas simples
9. Asegúrate de que el JSON sea válido y parseable


"""

    return prompt

def generar_documento_word(partida, asignatura, carrera, unidades_data):
    """Generar documento Word con formato simple y limpio"""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    import os
    import logging

    logger = logging.getLogger(__name__)

    try:
        doc = Document()

        # ---- Estilos propios (evitan azules de Heading 1) ----
        styles = doc.styles
        if 'UnidadTitulo' not in styles:
            unidad_style = styles.add_style('UnidadTitulo', WD_STYLE_TYPE.PARAGRAPH)
            unidad_style.font.name = 'Calibri'
            unidad_style.font.size = Pt(16)
            unidad_style.font.bold = True
            unidad_style.font.color.rgb = RGBColor(0, 0, 0)

        if 'PreguntaTitulo' not in styles:
            preg_style = styles.add_style('PreguntaTitulo', WD_STYLE_TYPE.PARAGRAPH)
            preg_style.font.name = 'Calibri'
            preg_style.font.size = Pt(12)
            preg_style.font.bold = True
            preg_style.font.color.rgb = RGBColor(0, 0, 0)

        # Configurar márgenes
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Crear encabezado solo con logo UNEMI
        def add_header_with_logo(doc_obj):
            section = doc_obj.sections[0]
            header = section.header

            # Solo agregar el logo en el encabezado
            logo_path = os.path.join('app', 'static', 'img', 'unemi.png')
            if os.path.exists(logo_path):
                try:
                    logo_paragraph = header.paragraphs[0]
                    logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run = logo_paragraph.add_run()
                    run.add_picture(logo_path, width=Inches(1.5))
                except Exception as e:
                    logger.warning(f"No se pudo cargar el logo: {e}")
                    # Si no se puede cargar la imagen, usar texto
                    logo_paragraph = header.paragraphs[0]
                    logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    logo_run = logo_paragraph.add_run('UNEMI')
                    logo_run.bold = True
                    logo_run.font.size = Pt(16)
            else:
                # Usar texto si no hay logo
                logo_paragraph = header.paragraphs[0]
                logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                logo_run = logo_paragraph.add_run('UNEMI')
                logo_run.bold = True
                logo_run.font.size = Pt(16)

        # Agregar encabezado
        add_header_with_logo(doc)

        # Agregar tabla de información de carrera y asignatura
        def add_info_table(doc_obj, carrera_obj, asignatura_obj):
            from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
            from docx.shared import Cm

            try:
                # Crear tabla con 2 filas y 2 columnas
                table = doc_obj.add_table(rows=2, cols=2)
                table.style = 'Table Grid'

                # Configurar propiedades de la tabla
                table.alignment = WD_TABLE_ALIGNMENT.LEFT
                table.left_indent = Cm(0)  # Sangría izquierda: 0 cm

                # Configurar ancho de columnas (12.356 cm total)
                table.columns[0].width = Cm(2.0)      # Columna izquierda (CARRERA, ASIGNATURA)
                table.columns[1].width = Cm(10.356)   # Columna derecha (valores)

                # Configurar altura mínima de filas (0.749 cm)
                for row in table.rows:
                    row.height = Cm(0.749)
                    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST  # Altura mínima

                # Configurar propiedades de celdas
                for row in table.rows:
                    for cell in row.cells:
                        # Alineación vertical: Parte superior
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

                # Primera fila - CARRERA
                row1 = table.rows[0]
                row1.cells[0].text = 'CARRERA'
                carrera_text = carrera_obj.get('descripcion', carrera_obj.get('nombre', '')) if carrera_obj else 'No especificada'
                row1.cells[1].text = str(carrera_text)[:100]  # Limitar longitud

                # Segunda fila - ASIGNATURA
                row2 = table.rows[1]
                row2.cells[0].text = 'ASIGNATURA'
                asignatura_text = asignatura_obj.get('descripcion', asignatura_obj.get('nombre', '')) if asignatura_obj else 'No especificada'
                row2.cells[1].text = str(asignatura_text)[:100]  # Limitar longitud

                # Aplicar formato a todas las celdas
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
                                run.font.size = Pt(11)
                                run.font.color.rgb = RGBColor(0, 0, 0)
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

                # Agregar espacio después de la tabla
                doc_obj.add_paragraph()

            except Exception as e:
                logger.error(f"Error al crear tabla de información: {e}")
                # Agregar información simple si falla la tabla
                doc_obj.add_paragraph('Información del Documento', style='UnidadTitulo')
                carrera_text = carrera_obj.get('descripcion', carrera_obj.get('nombre', '')) if carrera_obj else 'No especificada'
                p1 = doc_obj.add_paragraph(f"Carrera: {str(carrera_text)[:100]}")
                p1.runs[0].font.color.rgb = RGBColor(0, 0, 0)
                asignatura_text = asignatura_obj.get('descripcion', asignatura_obj.get('nombre', '')) if asignatura_obj else 'No especificada'
                p2 = doc_obj.add_paragraph(f"Asignatura: {str(asignatura_text)[:100]}")
                p2.runs[0].font.color.rgb = RGBColor(0, 0, 0)
                doc_obj.add_paragraph()

        # Agregar tabla de información
        add_info_table(doc, carrera, asignatura)

        # Contador global de preguntas para numeración continua
        contador_pregunta_global = 1

        # Preguntas por unidad
        for unidad in unidades_data:
            try:
                # Título de la unidad (sin color de tema)
                unidad_desc = str(unidad.get('descripcion', ''))[:200]  # Limitar longitud
                unidad_heading = doc.add_paragraph(
                    f'UNIDAD {unidad.get("numero", "?")}: {unidad_desc}',
                    style='UnidadTitulo'
                )
                unidad_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

                # Preguntas de la unidad
                for pregunta in unidad.get('preguntas', []):
                    try:
                        # Título de la pregunta con estilo propio
                        numero_pregunta = pregunta.get('numero', contador_pregunta_global)
                        pregunta_paragraph = doc.add_paragraph(
                            f'Pregunta {numero_pregunta}',
                            style='PreguntaTitulo'
                        )
                        pregunta_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

                        # Enunciado de la pregunta
                        enunciado_paragraph = doc.add_paragraph()
                        enunciado_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        enunciado_text = str(pregunta.get('enunciado', ''))[:500]  # Limitar longitud
                        enunciado_run = enunciado_paragraph.add_run(enunciado_text)
                        enunciado_run.font.size = Pt(11)
                        enunciado_run.font.color.rgb = RGBColor(0, 0, 0)

                        # Opciones
                        opciones_heading = doc.add_paragraph()
                        opciones_heading_run = opciones_heading.add_run('Opciones:')
                        opciones_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        opciones_heading_run.font.color.rgb = RGBColor(0, 0, 0)

                        opciones = pregunta.get('opciones', [])
                        for i, opcion in enumerate(opciones[:10], 1):  # Limitar a 10 opciones
                            try:
                                opcion_paragraph = doc.add_paragraph()
                                opcion_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                opcion_paragraph.paragraph_format.left_indent = Inches(0.3)

                                letra = chr(96 + i)  # a, b, c, d...
                                opcion_text = str(opcion.get('texto', ''))[:300]  # Limitar longitud
                                opcion_text = f'{letra}) {opcion_text}'
                                opcion_run = opcion_paragraph.add_run(opcion_text)
                                opcion_run.font.size = Pt(11)
                                opcion_run.font.color.rgb = RGBColor(0, 0, 0)
                                # Importante: NO poner en negrita la opción correcta
                            except Exception as e:
                                logger.warning(f"Error al procesar opción {i}: {e}")
                                continue

                        # Respuesta correcta (solo aquí se resalta si quieres mantenerlo)
                        respuesta_paragraph = doc.add_paragraph()
                        respuesta_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        respuesta_paragraph.paragraph_format.left_indent = Inches(0.3)

                        # Encontrar la respuesta correcta
                        for i, opcion in enumerate(opciones[:10], 1):
                            if opcion.get('es_correcta', False):
                                letra_correcta = chr(96 + i)
                                respuesta_run = respuesta_paragraph.add_run(f'Respuesta correcta: {letra_correcta}')
                                respuesta_run.bold = True
                                respuesta_run.font.size = Pt(11)
                                respuesta_run.font.color.rgb = RGBColor(0, 0, 0)
                                break

                        # Explicación de la respuesta
                        explicacion = pregunta.get('explicacion', '').strip()
                        if explicacion:
                            explicacion_paragraph = doc.add_paragraph()
                            explicacion_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            explicacion_paragraph.paragraph_format.left_indent = Inches(0.3)
                            
                            explicacion_heading = explicacion_paragraph.add_run('Explicación: ')
                            explicacion_heading.bold = True
                            explicacion_heading.font.size = Pt(11)
                            explicacion_heading.font.color.rgb = RGBColor(0, 0, 0)
                            
                            explicacion_text = str(explicacion)[:800]  # Limitar longitud
                            explicacion_run = explicacion_paragraph.add_run(explicacion_text)
                            explicacion_run.font.size = Pt(11)
                            explicacion_run.font.color.rgb = RGBColor(0, 0, 0)

                        # Incrementar contador global
                        contador_pregunta_global += 1

                        # Espacio entre preguntas
                        doc.add_paragraph()

                    except Exception as e:
                        logger.warning(f"Error al procesar pregunta {contador_pregunta_global}: {e}")
                        contador_pregunta_global += 1
                        continue

            except Exception as e:
                logger.warning(f"Error al procesar unidad {unidad.get('numero', '?')}: {e}")
                continue

        logger.info(f"Documento generado con {contador_pregunta_global - 1} preguntas")
        return doc

    except Exception as e:
        logger.error(f"Error crítico al generar documento: {e}", exc_info=True)
        raise e


# ============================================================================
# VISTA PARA CREAR PARTIDAS COMPLETAS - SUPABASE
# ============================================================================

class CrearPartidaCompletaView(View):
    """Vista para crear partida completa con programa analítico y unidades - Supabase"""

    def post(self, request):
        try:
            # Obtener datos del formulario
            descripcion = request.POST.get('descripcion', '').strip()
            asignatura_id = request.POST.get('asignatura')
            titulo_programa = request.POST.get('titulo_programa', '').strip()
            contexto = request.POST.get('contexto', '').strip()
            num_unidades = int(request.POST.get('num_unidades', 0))
            preguntas_por_unidad = int(request.POST.get('preguntas_por_unidad', 10))

            # Validaciones básicas
            if not all([descripcion, asignatura_id, titulo_programa, contexto, num_unidades]):
                messages.error(request, 'Todos los campos son obligatorios')
                return redirect('partida_lista')

            # 1. Crear la partida
            partida = PartidaRepository.create(
                descripcion=descripcion,
                asignatura_id=int(asignatura_id)
            )

            if not partida:
                messages.error(request, 'Error al crear la partida')
                return redirect('partida_lista')

            # 2. Crear el programa analítico
            programa = ProgramaAnaliticoRepository.create(
                titulo=titulo_programa,
                contexto=contexto,
                asignatura_id=int(asignatura_id)
            )

            if not programa:
                messages.error(request, 'Error al crear el programa analítico')
                return redirect('partida_lista')

            # 3. Crear las unidades
            unidades_creadas = []
            for i in range(1, num_unidades + 1):
                descripcion_unidad = request.POST.get(f'unidad_{i}_descripcion', '').strip()
                if descripcion_unidad:
                    unidad = UnidadRepository.create(
                        numero_unidad=i,
                        descripcion=descripcion_unidad,
                        num_preguntas=preguntas_por_unidad,
                        programa_analitico_id=programa['linea_educativa_id']
                    )
                    if unidad:
                        unidades_creadas.append(unidad)

            # 4. Crear preguntas automáticas para cada unidad con numeración secuencial
            preguntas_creadas = 0
            contador_pregunta_global = 1  # Contador global para numeración secuencial

            for unidad in unidades_creadas:
                for _ in range(preguntas_por_unidad):
                    pregunta = PreguntaRepository.create(
                        numero=contador_pregunta_global,  # Usar contador global
                        enunciado=f"Pregunta {contador_pregunta_global}",
                        unidad_id=unidad['unidad_id']
                    )
                    if pregunta:
                        # Crear opciones para la pregunta
                        for k in range(1, 5):  # 4 opciones por pregunta
                            es_correcta = (k == 1)  # La primera opción es correcta por defecto
                            OpcionRepository.create(
                                opcion=f"Opción {k} para pregunta {contador_pregunta_global}",
                                es_correcta=es_correcta,
                                pregunta_id=pregunta['pregunta_id']
                            )
                        preguntas_creadas += 1
                        contador_pregunta_global += 1  # Incrementar contador global

            # Mensaje de éxito
            messages.success(
                request,
                f'¡Partida creada exitosamente! '
                f'Se crearon {len(unidades_creadas)} unidades y {preguntas_creadas} preguntas automáticamente.'
            )

            return redirect('partida_lista')

        except Exception as e:
            messages.error(request, f'Error al crear la partida: {str(e)}')
            return redirect('partida_lista')
